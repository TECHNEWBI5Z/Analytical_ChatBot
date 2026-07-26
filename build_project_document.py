"""Build the Brinjal Insights project documentation DOCX."""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).parent
OUTPUT = ROOT / "Brinjal_Insights_Project_Documentation.docx"
IMAGE = ROOT / "assets" / "eggplant-varieties.jpg"

PLUM = "4B1C58"
DEEP_PLUM = "30173D"
LILAC = "F1E8F4"
GREEN = "356B48"
INK = "202124"
MUTED = "5F6368"
PALE_GREEN = "EDF5EC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    tc_pr.append(shade)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for column, width in zip(grid.gridCol_lst, widths):
        column.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = OxmlElement("w:tcMar")
            for side in ("top", "start", "bottom", "end"):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), "100" if side in ("start", "end") else "80")
                node.set(qn("w:type"), "dxa")
                margins.append(node)
            tc_pr.append(margins)


def font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_paragraph(doc, text="", style=None, before=0, after=6, color=INK, size=11, bold=False, italic=False):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.10
    if text:
        run = paragraph.add_run(text)
        font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.10
    for run in paragraph.runs:
        font(run, size=11)
    run = paragraph.add_run(text)
    font(run, size=11)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    font(run, size=16 if level == 1 else 13, color=PLUM if level == 1 else DEEP_PLUM, bold=True)
    return paragraph


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    for cell, value in zip(header.cells, headers):
        set_cell_shading(cell, LILAC)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        font(r, size=10, color=DEEP_PLUM, bold=True)
    for row_values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, row_values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            font(r, size=9.5)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in (("Heading 1", 16, PLUM), ("Heading 2", 13, DEEP_PLUM)):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    r = footer.add_run("Brinjal Insights | Local plant-breeding analytics prototype")
    font(r, size=8.5, color=MUTED)


def build_document():
    doc = Document()
    configure_styles(doc)

    # Editorial cover with a restrained plant-breeding identity.
    add_paragraph(doc, "PROJECT DOCUMENTATION", before=10, after=8, color=GREEN, size=10, bold=True)
    title = add_paragraph(doc, before=0, after=5)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    font(title.add_run("Brinjal Insights"), size=28, color=DEEP_PLUM, bold=True)
    subtitle = add_paragraph(doc, "A local AI analytics chatbot for eggplant plant-breeding data", after=14, color=MUTED, size=14)
    if IMAGE.exists():
        doc.add_picture(str(IMAGE), width=Inches(6.5))
        cap = add_paragraph(doc, "Eggplant varieties used as a visual identity for the prototype.", after=16, color=MUTED, size=8.5, italic=True)
    add_paragraph(doc, "Purpose", after=3, color=GREEN, size=10, bold=True)
    add_paragraph(
        doc,
        "Provide breeders and research teams with a simple, secure way to explore local physiological, root, and morphology measurements using text or voice questions—without paid LLM APIs.",
        after=14,
        size=11.5,
    )
    meta = add_table(
        doc,
        ["Project scope", "Deployment model", "Primary users"],
        [["Eggplant genotype data exploration", "Local machine / local MySQL", "Plant breeders, researchers, students"]],
        [3120, 3120, 3120],
    )
    add_paragraph(doc, "Version 1.0 | July 2026", after=0, color=MUTED, size=9)
    doc.add_page_break()

    add_heading(doc, "1. Problem statement")
    add_paragraph(
        doc,
        "Plant-breeding datasets often contain repeated genotype measurements across physiological, root, and morphological traits. Accessing these data usually requires SQL skills, spreadsheet manipulation, or repeated support from an analyst. This slows exploratory analysis and makes it harder for non-technical users to compare genotypes, calculate trait summaries, or examine relationships among measurements.",
    )
    add_paragraph(
        doc,
        "Brinjal Insights addresses this gap with a local, read-only conversational interface over MySQL. A user asks a question in natural language or by voice; the system generates a constrained MySQL SELECT query, validates it, retrieves the result, and presents a table plus a suitable visualization.",
    )

    add_heading(doc, "2. Data scope")
    add_paragraph(doc, "The prototype is designed around five imported CSV datasets, each containing 150 genotype records and linked conceptually through GenotypeName.")
    add_table(
        doc,
        ["Dataset", "Trait domain", "Example variables"],
        [
            ["phy_fst", "First physiological measurement", "CanopyTemp, SPAD, MSI, RWC, EpicuticularWax"],
            ["Phy_snd", "Second physiological measurement", "CanopyTemp, SPAD, MSI, RWC, EpicuticularWax"],
            ["Root_fst", "First root measurement", "RootLength, No_Lateral_Roots, R_S_Ratio"],
            ["Root_scd", "Second root measurement", "RootLength, No_Lateral_Roots, R_S_Ratio"],
            ["plant_morpho", "Morphology and fruit descriptors", "FruitShape, FruitColour, FlowerColour, StemPubescence"],
        ],
        [1700, 2600, 5060],
    )
    add_paragraph(doc, "Important modelling assumption: GenotypeName is the expected relationship key for comparisons and joins. The model is instructed to use only columns obtained directly from the MySQL schema.", color=GREEN, italic=True)

    add_heading(doc, "3. Solution overview")
    add_paragraph(doc, "The application uses a local-first architecture. It does not require paid cloud LLM APIs, and the MySQL account is intentionally limited to SELECT permission.")
    add_table(
        doc,
        ["Stage", "What happens", "Control"],
        [
            ["1. User question", "User types a question or records it in the browser.", "One question per request."],
            ["2. Speech-to-text", "Whisper transcribes an English recording locally when voice is used.", "No paid speech API."],
            ["3. SQL generation", "Ollama produces one MySQL SELECT statement from the live database schema.", "ART prompt and temperature 0."],
            ["4. SQL safety", "The application rejects multi-statement and non-read-only SQL, then applies a row cap.", "sqlglot parsing plus forbidden-command checks."],
            ["5. Analysis output", "Pandas retrieves the result and Streamlit shows the table, metric, or graph.", "Interactive chart controls."],
        ],
        [1550, 4700, 3110],
    )

    add_heading(doc, "4. Core features")
    add_table(
        doc,
        ["Feature", "User value"],
        [
            ["Natural-language querying", "Ask questions such as ‘What is the average RootLength grouped by FruitShape?’ without writing SQL."],
            ["Voice input", "Record an English question; transcription starts automatically and sends the result to the same analytics workflow."],
            ["Schema-aware querying", "The live MySQL schema is provided to the model so it can select existing tables and fields."],
            ["Joins and aggregates", "Supports INNER JOIN, AVG, COUNT, MIN, MAX, GROUP BY, HAVING, and ORDER BY for genotype-level analysis."],
            ["Interactive visualization", "Bar, line, scatter, and frequency charts are offered from the returned data."],
            ["Read-only safeguards", "The MySQL user has SELECT-only access; application validation rejects dangerous operations."],
            ["Local AI operation", "Ollama and Whisper run locally, avoiding paid API requirements and reducing data exposure."],
        ],
        [2500, 6860],
    )

    add_heading(doc, "5. Methodology")
    add_heading(doc, "5.1 Prompt methodology: ART", level=2)
    add_paragraph(doc, "The SQL-generation system prompt follows ART: Action, Role, and Task.")
    add_bullet(doc, "Action: translate one plant-breeding analytics question into one read-only MySQL query.")
    add_bullet(doc, "Role: act as a careful MySQL analyst for eggplant breeding experiments, prioritizing correct identifiers, joins, and aggregate aliases.")
    add_bullet(doc, "Task: return JSON containing the SQL and a short explanation, while following MySQL syntax and safety requirements.")

    add_heading(doc, "5.2 Query execution methodology", level=2)
    add_bullet(doc, "Discover tables and columns at runtime using SQLAlchemy inspection; no hard-coded database schema is required.")
    add_bullet(doc, "Use a deterministic shortcut for common ‘show first N rows from table’ requests to avoid unnecessary model errors.")
    add_bullet(doc, "Parse generated SQL with sqlglot, reject write operations and multiple statements, and limit returned rows to 200 by default.")
    add_bullet(doc, "Execute only after validation, using a MySQL account created specifically with SELECT permission.")

    add_heading(doc, "5.3 Visualization methodology", level=2)
    add_paragraph(doc, "The result shape determines the available presentation. A single-row aggregation is displayed as a metric card. Categorical plus numeric results can be shown as bar or line charts. Two numeric fields can be compared as a scatter plot. Categorical results can be transformed into a frequency graph.")

    add_heading(doc, "6. Models, libraries, and add-ins")
    add_table(
        doc,
        ["Component", "Technology", "Purpose"],
        [
            ["Local LLM", "Ollama with qwen2.5:3b", "Converts natural-language plant-breeding questions into MySQL SELECT queries."],
            ["Speech-to-text", "faster-whisper with tiny.en", "Locally transcribes English voice questions. base.en can be selected for higher accuracy."],
            ["Web application", "Streamlit", "Provides the chat interface, microphone recording, tables, charts, and themed UI."],
            ["Database layer", "SQLAlchemy + PyMySQL", "Connects to MySQL, reads schema metadata, and executes validated queries."],
            ["SQL guardrail", "sqlglot", "Parses and validates MySQL syntax before execution."],
            ["Data handling", "pandas", "Loads SQL results and prepares chart data."],
            ["Visualization", "Streamlit + Vega-Lite", "Renders bar, line, scatter, and frequency graphs."],
        ],
        [1900, 2850, 4610],
    )
    add_paragraph(doc, "No paid LLM API, OpenAI API key, or cloud database connector is required for the current implementation.", color=GREEN, bold=True)

    add_heading(doc, "7. Plant-breeding use cases")
    add_table(
        doc,
        ["Use case", "Example question", "Expected analysis"],
        [
            ["Root trait screening", "Which genotypes have the highest R_S_Ratio in Root_scd?", "Ranks genotypes by root-to-shoot ratio."],
            ["Physiology comparison", "Compare average SPAD in phy_fst and Phy_snd.", "Compares chlorophyll-related readings across measurements."],
            ["Trait association", "Show RootLength and SPAD for each GenotypeName.", "Joins root and physiological traits for cross-trait inspection."],
            ["Morphology grouping", "What is the average RootLength grouped by FruitShape?", "Relates root performance to fruit-shape classes."],
            ["Frequency screening", "Count plants by FruitColour.", "Summarizes categorical diversity in the breeding material."],
            ["Outlier exploration", "Show GenotypeName, RootLength, and R_S_Ratio from Root_fst.", "Supports scatter-plot review of potential outliers or selections."],
        ],
        [1900, 3700, 3760],
    )

    add_heading(doc, "8. Operating procedure")
    for step in (
        "Open the project folder in VS Code and activate the Python virtual environment.",
        "Ensure MySQL and the Ollama application are running.",
        "Start the app with `python -m streamlit run app.py`.",
        "Ask one plain-English question at a time, by typing or voice recording.",
        "Review the expandable SQL statement during early testing, then inspect the table and chart.",
        "If results appear incorrect, check the database schema in the sidebar and rephrase with exact table/trait names.",
    ):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(step)
        font(r, size=11)

    add_heading(doc, "9. Limitations and next improvements")
    add_bullet(doc, "Small local models can still select an incorrect valid column or propose an inappropriate join. Users should review generated SQL during validation.")
    add_bullet(doc, "The current voice model is English-focused. Multilingual use would require a multilingual Whisper model and testing with field terminology.")
    add_bullet(doc, "The current app keeps the latest answer in the session. Persistent conversation history, saved query bookmarks, and user authentication are future enhancements.")
    add_bullet(doc, "A breeding-specific data dictionary could define trial stages, environments, replications, units, and trait meanings to improve scientific interpretation.")

    add_heading(doc, "10. Security and governance")
    add_paragraph(doc, "The design uses defense in depth. The database account must have SELECT-only access. The application rejects write-oriented SQL and multi-statement input, parses SQL before execution, and limits results. These safeguards reduce risk but do not replace normal database administration, backups, access control, or scientific data-quality review.")

    add_heading(doc, "Appendix: example SQL pattern")
    example = add_paragraph(doc, before=2, after=4)
    example.paragraph_format.left_indent = Inches(0.2)
    example.paragraph_format.right_indent = Inches(0.2)
    for line in (
        "SELECT rf.GenotypeName, rf.RootLength, pf.SPAD",
        "FROM Root_fst AS rf",
        "INNER JOIN phy_fst AS pf ON rf.GenotypeName = pf.GenotypeName",
        "ORDER BY rf.RootLength DESC",
        "LIMIT 200;",
    ):
        run = example.add_run(line + "\n")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(DEEP_PLUM)

    add_paragraph(doc, "Image credit: J.E. Fee, ‘Three Types of Eggplant,’ Wikimedia Commons, CC BY 2.0.", after=0, color=MUTED, size=8.5)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
